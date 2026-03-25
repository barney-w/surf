"""DOCX connector for extracting text and creating IngestedDocument instances."""

from __future__ import annotations

import hashlib
import re
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from pathlib import Path

from docx import Document  # pyright: ignore[reportMissingTypeStubs]

from src.models import DocumentMetadata, IngestedDocument

_SAFE_URL_RE = re.compile(r"^https?://", re.IGNORECASE)

# Mapping from python-docx style names to Markdown-style heading prefixes.
_HEADING_PREFIX: dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
}


def _validate_source_url(url: str | None) -> str | None:
    if url is None:
        return None
    if not _SAFE_URL_RE.match(url):
        raise ValueError(f"Invalid source_url: must start with http:// or https://, got: {url!r}")
    return url


def _table_to_text(table: Any) -> str:
    """Convert a python-docx Table object to a pipe-delimited text representation."""
    rows: list[str] = []
    for row in table.rows:  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        cells = [
            cell.text.strip()  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
            for cell in row.cells  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
        ]
        rows.append("| " + " | ".join(cells) + " |")  # pyright: ignore[reportUnknownArgumentType]
    return "\n".join(rows)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a DOCX file, preserving heading hierarchy and tables.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text with headings formatted as Markdown-style prefixes
        and tables rendered as pipe-delimited rows.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a DOCX file.
    """
    if not file_path.exists():
        msg = f"File not found: {file_path}"
        raise FileNotFoundError(msg)

    if file_path.suffix.lower() != ".docx":
        msg = f"Not a DOCX file: {file_path}"
        raise ValueError(msg)

    doc = Document(str(file_path))  # pyright: ignore[reportUnknownMemberType]

    # Build lookup maps keyed by the id() of each XML element so we can
    # iterate the body's children in document order while using the
    # fully-initialised Paragraph / Table proxy objects.
    para_map: dict[int, Any] = {
        id(p._element): p  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType, reportPrivateUsage]
        for p in doc.paragraphs  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
    }
    table_map: dict[int, Any] = {
        id(t._element): t  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType, reportPrivateUsage]
        for t in doc.tables  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
    }

    from docx.oxml.ns import qn  # pyright: ignore[reportMissingTypeStubs]

    blocks: list[str] = []

    for child in doc.element.body:  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType, reportOptionalMemberAccess]
        if child.tag == qn("w:p"):  # pyright: ignore[reportUnknownMemberType]
            para = para_map.get(id(child))  # pyright: ignore[reportUnknownArgumentType]
            if para is None:
                continue
            text = para.text.strip()  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
            if not text:
                continue
            style_name = para.style.name if para.style else ""  # pyright: ignore[reportUnknownMemberType, reportUnknownVariableType]
            prefix = _HEADING_PREFIX.get(style_name, "")  # pyright: ignore[reportUnknownArgumentType]
            blocks.append(f"{prefix}{text}")
        elif child.tag == qn("w:tbl"):  # pyright: ignore[reportUnknownMemberType]
            table_obj = table_map.get(id(child))  # pyright: ignore[reportUnknownArgumentType]
            if table_obj is not None:
                table_text = _table_to_text(table_obj)
                if table_text.strip():
                    blocks.append(table_text)

    return "\n\n".join(blocks)


def _generate_document_id(file_path: Path) -> str:
    """Generate a deterministic document ID from the file path."""
    return hashlib.sha256(f"docx:{file_path.name}".encode()).hexdigest()[:16]


def create_document_from_docx(file_path: Path, manifest_entry: dict[str, Any]) -> IngestedDocument:
    """Create an IngestedDocument from a DOCX file and its manifest entry.

    Args:
        file_path: Path to the DOCX file.
        manifest_entry: Dictionary with document metadata. Expected keys:
            - domain (str): Document domain, e.g. "hr", "it", "governance"
            - document_type (str): e.g. "policy", "procedure", "agreement"
            - title (str, optional): Document title. Defaults to file stem.
            - raw_path (str, optional): Blob storage path. Defaults to str(file_path).
            - content_source (str, optional): e.g. "website", "sharepoint". Defaults to "".
            - section_path (str, optional): URL/folder path. Defaults to "".
            - version, effective_date, expiry_date, author, source_url, tags (optional)

    Returns:
        An IngestedDocument populated with extracted text and metadata.
    """
    content = extract_text_from_docx(file_path)
    doc_id = _generate_document_id(file_path)

    metadata = DocumentMetadata(
        domain=manifest_entry["domain"],
        document_type=manifest_entry["document_type"],
        version=manifest_entry.get("version"),
        effective_date=manifest_entry.get("effective_date"),
        expiry_date=manifest_entry.get("expiry_date"),
        author=manifest_entry.get("author"),
        source_url=_validate_source_url(manifest_entry.get("source_url")),
        tags=manifest_entry.get("tags", []),
        content_source=manifest_entry.get("content_source", ""),
        section_path=manifest_entry.get("section_path", ""),
    )

    return IngestedDocument(
        id=doc_id,
        source="docx",
        title=manifest_entry.get("title", file_path.stem),
        content=content,
        metadata=metadata,
        raw_path=manifest_entry.get("raw_path", str(file_path)),
    )
